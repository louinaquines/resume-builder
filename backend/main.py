from fastapi import FastAPI, Response, Request
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import StreamingResponse
from pydantic import BaseModel
from dotenv import load_dotenv
import httpx
import os
import json
from reportlab.lib.pagesizes import letter
from reportlab.lib import colors
from reportlab.lib.units import inch
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, HRFlowable, Table, TableStyle
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.enums import TA_LEFT, TA_CENTER
from io import BytesIO
import base64
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

load_dotenv()

OPENROUTER_API_KEY = os.getenv("OPENROUTER_API_KEY")
OPENROUTER_URL = "https://openrouter.ai/api/v1/chat/completions"
MODEL = "openrouter/auto"

app = FastAPI()

app.add_middleware(
    CORSMiddleware,
    allow_origins=[
        "http://localhost:5173",
        "https://ai-resume-ph.vercel.app",
    ],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
    expose_headers=["Content-Disposition"],
)

class ResumeRequest(BaseModel):
    full_name: str = ""
    email: str = ""
    phone: str = ""
    location: str = ""
    linkedin: str = ""
    photo: str = ""
    social_link: str = ""
    headline: str = ""
    date_of_birth: str = ""
    civil_status: str = ""
    nationality: str = ""
    gender: str = ""
    job_title: str = ""
    job_description: str = ""
    job_type: str = "Full-time"
    availability: str = ""
    work_experience: list[dict] = []
    education: list[dict] = []
    certifications: list[dict] = []
    languages: list[dict] = []
    skills: str = ""
    tone: str = "Professional"
    references: list[dict] = []
    seminars: list[dict] = []
    template: str = "Minimalist"
    goal: str = "Applying for a Job"
    bullet_style: str = "Short & Punchy"
    color_accent: str = "Classic Black"

def build_prompt(data: ResumeRequest) -> str:
    work_block = ""
    for job in data.work_experience:
        work_block += f"""
Company: {job.get('company', '')}
Role: {job.get('role', '')}
Duration: {job.get('duration', '')}
Responsibilities: {job.get('responsibilities', '')}
"""

    edu_block = ""
    for edu in data.education:
        edu_block += f"""
School: {edu.get('school', '')}
Degree: {edu.get('degree', '')}
Year: {edu.get('year', '')}
"""

    cert_block = ""
    for cert in (data.certifications or []):
        cert_block += f"- {cert.get('name', '')} | {cert.get('issuer', '')} | {cert.get('year', '')}\n"

    lang_block = ""
    for lang in (data.languages or []):
        lang_block += f"- {lang.get('language', '')} ({lang.get('proficiency', '')})\n"

    ref_block = ""
    for ref in (data.references or []):
        ref_block += f"- {ref.get('name', '')} | {ref.get('position', '')} | {ref.get('company', '')} | {ref.get('contact', '')}\n"

    seminar_block = ""
    for sem in (data.seminars or []):
        seminar_block += f"- {sem.get('title', '')} | {sem.get('organizer', '')} | {sem.get('year', '')}\n"

    tone_instructions = {
        "Professional": """
- Use formal, polished language
- Focus on achievements, metrics, and measurable impact
- Keep sentences concise and direct
- Use strong action verbs: Managed, Led, Achieved, Delivered, Implemented
- Avoid casual language or personal pronouns
""",
        "Modern": """
- Use clean, contemporary language that feels fresh and confident
- Balance professionalism with personality
- Highlight adaptability, collaboration, and innovation
- Use active voice and dynamic verbs: Drove, Built, Launched, Scaled, Optimized
- Keep it punchy — short impactful sentences
""",
        "Creative": """
- Use expressive, engaging language that shows personality
- Tell a story — make the candidate memorable
- Highlight passion, creativity, and unique perspective
- Use vivid verbs and descriptive language: Crafted, Designed, Pioneered, Transformed, Reimagined
- It should feel human and warm, not corporate
""",
        "Executive": """
- Use authoritative, strategic language befitting senior leadership
- Focus on business impact, revenue, team leadership, and organizational growth
- Emphasize vision, strategy, and high-level decision making
- Use commanding verbs: Spearheaded, Directed, Orchestrated, Championed, Transformed
- Quantify everything — percentages, team sizes, budgets, timelines
""",
    }

    selected_tone = tone_instructions.get(data.tone, tone_instructions["Professional"])

    return f"""You are a professional resume writer in the Philippines. Generate a complete, polished resume using ONLY the data provided below. Do not invent, assume, or add any information not explicitly given.

Tone Style: {data.tone}
Tone Instructions:
{selected_tone}

Target Job Title: {data.job_title}
Job Type: {data.job_type}
Availability: {data.availability}

PERSONAL INFO:
Name: {data.full_name}
Headline: {data.headline}
Email: {data.email}
Phone: {data.phone}
Location: {data.location}
LinkedIn: {data.linkedin}
Date of Birth: {data.date_of_birth}
Civil Status: {data.civil_status}
Nationality: {data.nationality}

WORK EXPERIENCE:
{work_block if work_block.strip() else "None provided"}

EDUCATION:
{edu_block if edu_block.strip() else "None provided"}

CERTIFICATIONS & LICENSES:
{cert_block.strip() if cert_block.strip() else "None provided"}

LANGUAGES:
{lang_block.strip() if lang_block.strip() else "None provided"}

SKILLS:
{data.skills if data.skills.strip() else "None provided"}

SEMINARS & TRAININGS:
{seminar_block.strip() if seminar_block.strip() else "None provided"}

CHARACTER REFERENCES:
{ref_block.strip() if ref_block.strip() else "None provided"}

STRICT RULES — YOU MUST FOLLOW THESE:
1. Output ONLY real data from the sections above. Never invent or assume anything.
2. If a section says "None provided", output only the section header — NO content, NO dashes, NO placeholder text.
3. NEVER use square brackets like [Name], [Organization], [Date], [Relevant], [Year Obtained] anywhere.
4. NEVER use ++ or == markup around text.
5. Use ** only for bold category labels in the Skills section.
6. Write a 3-4 sentence professional summary using only the actual data provided.
7. Rewrite work experience bullets to be achievement-focused using only the actual responsibilities given.

Output in this exact markdown structure:

# {data.full_name}
{data.email} | {data.phone} | {data.location}{f" | {data.linkedin}" if data.linkedin else ""}

## Personal Information
**Date of Birth:** {data.date_of_birth if data.date_of_birth else "N/A"} | **Civil Status:** {data.civil_status if data.civil_status else "N/A"} | **Nationality:** {data.nationality if data.nationality else "N/A"}

## Professional Summary
[write 3-4 sentences here based ONLY on actual data above]

## Work Experience
[only if work experience was provided, otherwise leave blank under this header]

## Education
[only if education was provided, otherwise leave blank under this header]

## Certifications & Licenses
[only if certifications were provided, otherwise leave blank under this header]

## Skills
[only if skills were provided, otherwise leave blank under this header]

## Languages
[only if languages were provided, otherwise leave blank under this header]

## Seminars & Trainings
[only if seminars were provided, otherwise leave blank under this header]

## Character References
[only if references were provided, otherwise write "Available upon request."]

Output only the resume. No commentary. No placeholder brackets. No invented data.
"""

@app.post("/debug")
async def debug(request: Request):
    body = await request.json()
    return body

@app.post("/generate")
async def generate_resume(data: ResumeRequest):
    prompt = build_prompt(data)

    async def stream():
        try:
            async with httpx.AsyncClient(timeout=60) as client:
                async with client.stream(
                    "POST",
                    OPENROUTER_URL,
                    headers={
                        "Authorization": f"Bearer {OPENROUTER_API_KEY}",
                        "Content-Type": "application/json",
                        "HTTP-Referer": "https://ai-resume-ph.vercel.app",
                        "X-Title": "AI Resume Builder",
                    },
                    json={
                        "model": MODEL,
                        "stream": True,
                        "messages": [
                            {"role": "user", "content": prompt}
                        ],
                    },
                ) as response:
                    async for line in response.aiter_lines():
                        if line.startswith("data: "):
                            data_str = line[6:]
                            if data_str.strip() == "[DONE]":
                                break
                            try:
                                chunk = json.loads(data_str)
                                delta = chunk["choices"][0]["delta"].get("content", "")
                                if delta:
                                    yield delta
                            except Exception:
                                continue
        except Exception as e:
            yield f"ERROR: {str(e)}"

    return StreamingResponse(stream(), media_type="text/plain")

@app.post("/generate-pdf")
async def generate_pdf(data: ResumeRequest):
    prompt = build_prompt(data)

    async with httpx.AsyncClient(timeout=60) as client:
        response = await client.post(
            OPENROUTER_URL,
            headers={
                "Authorization": f"Bearer {OPENROUTER_API_KEY}",
                "Content-Type": "application/json",
                "HTTP-Referer": "https://ai-resume-ph.vercel.app",
                "X-Title": "AI Resume Builder",
            },
            json={
                "model": MODEL,
                "stream": False,
                "messages": [{"role": "user", "content": prompt}],
            },
        )
        result = response.json()
        markdown_text = result["choices"][0]["message"]["content"]

    # Parse markdown
    def parse_section(md, section_name):
        lines = md.split("\n")
        result = []
        in_section = False
        for line in lines:
            if line.strip().startswith("## ") and section_name.lower() in line.lower():
                in_section = True
                continue
            if in_section and line.strip().startswith("## "):
                break
            if in_section and line.strip():
                result.append(line.strip())
        return result

    summary_lines = parse_section(markdown_text, "summary") or parse_section(markdown_text, "objective")
    summary_text = " ".join(l for l in summary_lines if not l.startswith("#"))

    # Color
    accent_map = {"Deep Purple": "#4c1d95", "Electric Blue": "#2563eb", "Classic Black": "#1e293b", "Soft Green": "#059669"}
    tone_color_map = {"Professional": "#1e293b", "Modern": "#2563eb", "Creative": "#7c3aed", "Executive": "#1e3a5f"}
    if data.color_accent and data.color_accent != "Classic Black":
        hex_color = accent_map.get(data.color_accent, "#1e293b")
    else:
        hex_color = tone_color_map.get(data.tone, "#1e293b")

    dark = colors.HexColor(hex_color)
    gray_bg = colors.HexColor("#f3f4f6")
    text_dark = colors.HexColor("#1e293b")
    text_mid = colors.HexColor("#4b5563")
    text_light = colors.HexColor("#9ca3af")
    line_color = colors.HexColor("#e5e7eb")
    white = colors.white

    buffer = BytesIO()
    PAGE_W, PAGE_H = letter
    SIDEBAR_W = 2.4 * inch
    MAIN_W = PAGE_W - SIDEBAR_W
    HEADER_H = 1.2 * inch

    doc = SimpleDocTemplate(buffer, pagesize=letter,
        leftMargin=0, rightMargin=0, topMargin=0, bottomMargin=0)

    styles = getSampleStyleSheet()

    def S(name, **kw): return ParagraphStyle(name, **kw)

    name_s = S("N", fontName="Helvetica-Bold", fontSize=18, textColor=white, leading=22, spaceAfter=2)
    role_s = S("R", fontName="Helvetica", fontSize=11, textColor=colors.HexColor("#cbd5e1"), leading=14)
    sidebar_title_s = S("ST", fontName="Helvetica-Bold", fontSize=8, textColor=text_dark,
                         leading=10, spaceAfter=2, spaceBefore=12)
    sidebar_text_s = S("Stx", fontName="Helvetica", fontSize=8, textColor=text_mid,
                        leading=11, spaceAfter=2, wordWrap='CJK')
    main_title_s = S("MT", fontName="Helvetica-Bold", fontSize=10, textColor=text_dark,
                      leading=13, spaceAfter=4, spaceBefore=10)
    main_body_s = S("MB", fontName="Helvetica", fontSize=8, textColor=text_mid,
                     leading=12, spaceAfter=2)
    job_title_s = S("JT", fontName="Helvetica-Bold", fontSize=9, textColor=text_dark,
                     leading=12, spaceAfter=2)
    empty_s = S("E", fontName="Helvetica-Oblique", fontSize=8, textColor=text_light, leading=11)

    def sidebar_block(title, items, empty_text="—"):
        out = [
            Paragraph(title.upper(), sidebar_title_s),
            HRFlowable(width=SIDEBAR_W - 0.4*inch, thickness=0.5, color=colors.HexColor("#d1d5db"), spaceAfter=4),
        ]
        if items:
            for item in items:
                out.append(Paragraph(str(item), sidebar_text_s))
        else:
            out.append(Paragraph(empty_text, empty_s))
        return out

    def main_block(title, items, empty_text="—"):
        out = [
            Paragraph(title.upper(), main_title_s),
            HRFlowable(width=MAIN_W - 0.5*inch, thickness=1, color=text_dark, spaceAfter=6),
        ]
        if items:
            for item in items:
                out.append(item)
        else:
            out.append(Paragraph(empty_text, empty_s))
        out.append(Spacer(1, 6))
        return out

    # Build contact
    contact_items = [c for c in [data.email, data.phone, data.location, data.linkedin] if c]

    personal_items = [p for p in [
        f"DOB: {data.date_of_birth}" if data.date_of_birth else "",
        f"Gender: {data.gender}" if data.gender else "",
        f"Civil Status: {data.civil_status}" if data.civil_status else "",
        f"Nationality: {data.nationality}" if data.nationality else "",
        f"Available: {data.availability}" if data.availability else "",
        f"Job Type: {data.job_type}" if data.job_type else "",
    ] if p]

    skill_items = [s.strip() for s in data.skills.split(",") if s.strip()]

    cert_items = [
        f"{c.get('name','')}{' — ' + c.get('issuer','') if c.get('issuer') else ''}{' (' + c.get('year','') + ')' if c.get('year') else ''}"
        for c in data.certifications if c.get("name")
    ]

    lang_items = [
        f"{l.get('language','')}{' (' + l.get('proficiency','') + ')' if l.get('proficiency') else ''}"
        for l in data.languages if l.get("language")
    ]

    # Build experience items
    exp_flowables = []
    for j in data.work_experience:
        if not (j.get("company") or j.get("role")): continue
        parts = []
        if j.get("role"): parts.append(j["role"])
        if j.get("company"): parts.append(j["company"])
        dur = f" | {j['duration']}" if j.get("duration") else ""
        exp_flowables.append(Paragraph(f"<b>{' — '.join(parts)}{dur}</b>", job_title_s))
        for b in (j.get("responsibilities") or "").split("\n"):
            b = b.strip()
            if b:
                exp_flowables.append(Paragraph(f"• {b}", main_body_s))
        exp_flowables.append(Spacer(1, 4))

    edu_flowables = []
    for e in data.education:
        if not (e.get("school") or e.get("degree")): continue
        text = f"<b>{e.get('degree','')}</b>{' — ' + e.get('school','') if e.get('school') else ''}{' | ' + e.get('year','') if e.get('year') else ''}"
        edu_flowables.append(Paragraph(text, main_body_s))
        edu_flowables.append(Spacer(1, 3))

    seminar_flowables = [
        Paragraph(f"• {s.get('title','')}{' — ' + s.get('organizer','') if s.get('organizer') else ''}{' (' + s.get('year','') + ')' if s.get('year') else ''}", main_body_s)
        for s in data.seminars if s.get("title")
    ]

    ref_flowables = [
        Paragraph(f"• {' | '.join(filter(None, [r.get('name'), r.get('position'), r.get('company'), r.get('contact')]))}", main_body_s)
        for r in data.references if r.get("name")
    ] or [Paragraph("Available upon request.", main_body_s)]

    # Assemble sidebar
    sidebar_cells = []
    sidebar_cells += sidebar_block("Contact", contact_items)
    sidebar_cells += sidebar_block("Personal Info", personal_items)
    sidebar_cells += sidebar_block("Key Skills", skill_items)
    sidebar_cells += sidebar_block("Certifications", cert_items)
    sidebar_cells += sidebar_block("Languages", lang_items)

    # Assemble main
    main_cells = []
    main_cells += main_block("Career Objective",
        [Paragraph(summary_text, main_body_s)] if summary_text else [], "—")
    main_cells += main_block("Work Experience", exp_flowables)
    main_cells += main_block("Education", edu_flowables)
    main_cells += main_block("Seminars & Trainings", seminar_flowables)
    main_cells += main_block("Character References", ref_flowables)

    # Header
    photo_cell = Paragraph("", styles["Normal"])
    if data.photo and data.photo.startswith("data:image"):
        try:
            from reportlab.platypus import Image as RLImage
            img_bytes = base64.b64decode(data.photo.split(",", 1)[1])
            photo_cell = RLImage(BytesIO(img_bytes), width=0.75*inch, height=0.75*inch)
        except Exception:
            pass

    header_table = Table(
        [[photo_cell, [Paragraph(data.full_name.upper(), name_s), Paragraph(data.job_title or data.headline or "", role_s)]]],
        colWidths=[1.0*inch, PAGE_W - 1.0*inch],
        rowHeights=[HEADER_H],
    )
    header_table.setStyle(TableStyle([
        ("BACKGROUND", (0, 0), (-1, -1), dark),
        ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
        ("LEFTPADDING", (0, 0), (0, 0), 20),
        ("LEFTPADDING", (1, 0), (1, 0), 14),
        ("TOPPADDING", (0, 0), (-1, -1), 14),
        ("BOTTOMPADDING", (0, 0), (-1, -1), 14),
    ]))

    body_table = Table(
        [[sidebar_cells, main_cells]],
        colWidths=[SIDEBAR_W, MAIN_W],
    )
    body_table.setStyle(TableStyle([
        ("BACKGROUND", (0, 0), (0, 0), gray_bg),
        ("VALIGN", (0, 0), (-1, -1), "TOP"),
        ("LEFTPADDING", (0, 0), (0, 0), 12),
        ("RIGHTPADDING", (0, 0), (0, 0), 12),
        ("TOPPADDING", (0, 0), (-1, -1), 12),
        ("LEFTPADDING", (1, 0), (1, 0), 16),
        ("RIGHTPADDING", (1, 0), (1, 0), 16),
    ]))

    doc.build([header_table, body_table])
    return Response(
        content=buffer.getvalue(),
        media_type="application/pdf",
        headers={"Content-Disposition": "attachment; filename=resume.pdf"}
    )
@app.post("/chat")
async def chat(request: Request):
    body = await request.json()
    messages = body.get("messages", [])
    
    system_prompt = """You are a friendly, conversational AI resume builder.

CRITICAL OUTPUT CONSTRAINTS:
1. MAX ONE QUESTION PER MESSAGE. You must NEVER ask for more than one piece of information in a single response.
2. KEEP IT SHORT & GENTLE. Your response should never be longer than 2 to 3 sentences. Do not overwhelm the user.
3. ONE AT A TIME: If you ask for an email, STOP. Do not ask for their location or anything else until they reply.
4. BE LENIENT: If they say "none", "skip", or "same", just say "No problem!" and move to the very next single item. Do not repeat questions.
5. NEVER simulate the user's side of the conversation.

INFORMATION CHECKLIST (You must ask exactly one step at a time in this strict order):

Step 1: Ask for Full Name and Professional Title. (User types)
Step 2: Ask for Resume Tone.
Step 3: Ask for Email, Phone, and LinkedIn. (User types)
Step 4: Ask for Template style.
Step 5: Ask for recent work experience and achievements. (User types)
Step 6: Ask for the Goal of the resume.
Step 7: Ask for Software Tools and Skills. (User types)
Step 8: Ask for Bullet Style.
Step 9: Ask for Education and Degree. (User types)
Step 10: Ask for Color Accent.

HOW TO START AND PROGRESS:
- You MUST start every single message with `STEP X:` where X is the number of the step you are on.
- FIRST MESSAGE ONLY: Output `STEP 1:` followed by greeting the user and asking for their name and title.
- FOLLOWING MESSAGES: Output `STEP X:` followed by acknowledging their previous answer and asking the next question.

CRITICAL FORMATTING RULES:
- Never output brackets or internal thoughts.
- Start every message with `STEP X:`
- YOU MUST STOP and wait for the user between steps. NEVER output multiple steps in one message.

ONLY when you have completed all 10 steps, output EXACTLY this JSON block and nothing else:

RESUME_READY:{"full_name":"...","headline":"...","tone":"...","email":"...","phone":"...","linkedin":"...","social_link":"...","template":"...","work_experience":[{"company":"...","role":"...","duration":"...","responsibilities":"..."}],"goal":"...","skills":"...","bullet_style":"...","education":[{"school":"...","degree":"...","year":"..."}],"color_accent":"..."}"""

    async def stream():
        try:
            async with httpx.AsyncClient(timeout=60) as client:
                payload_messages = [
                    {"role": "system", "content": system_prompt},
                    *messages,
                    {"role": "system", "content": "CRITICAL REMINDER: Max 2 sentences. DO NOT write internal thoughts. Start your message with `STEP X:` indicating the current step number. Then ask the question and IMMEDIATELY STOP."}
                ]
                
                async with client.stream(
                    "POST",
                    OPENROUTER_URL,
                    headers={
                        "Authorization": f"Bearer {OPENROUTER_API_KEY}",
                        "Content-Type": "application/json",
                        "HTTP-Referer": "hhttps://ai-resume-ph.vercel.app",
                        "X-Title": "AI Resume Builder",
                    },
                    json={
                        "model": MODEL,
                        "stream": True,
                        "stop": ["Student:", "User:", "\nStudent", "\nUser", "user:", "student:", "\nSTEP", "STEP X"],
                        "messages": payload_messages,
                    },
                ) as response:
                    async for line in response.aiter_lines():
                        if line.startswith("data: "):
                            data_str = line[6:]
                            if data_str.strip() == "[DONE]":
                                break
                            try:
                                chunk = json.loads(data_str)
                                delta = chunk["choices"][0]["delta"].get("content", "")
                                if delta:
                                    yield delta
                            except Exception:
                                continue
        except Exception as e:
            yield f"ERROR: {str(e)}"

    return StreamingResponse(stream(), media_type="text/plain")

@app.post("/generate-docx")
async def generate_docx(data: ResumeRequest):
    doc = Document()

    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    dark = RGBColor(0x1e, 0x29, 0x3b)
    mid = RGBColor(0x4b, 0x55, 0x63)
    light = RGBColor(0x9c, 0xa3, 0xaf)

    def heading(text, size=11):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(14)
        p.paragraph_format.space_after = Pt(2)
        run = p.add_run(text.upper())
        run.bold = True
        run.font.size = Pt(size)
        run.font.color.rgb = dark
        # Underline via border
        from docx.oxml.ns import qn
        from docx.oxml import OxmlElement
        pPr = p._p.get_or_add_pPr()
        pBdr = OxmlElement('w:pBdr')
        bottom = OxmlElement('w:bottom')
        bottom.set(qn('w:val'), 'single')
        bottom.set(qn('w:sz'), '6')
        bottom.set(qn('w:space'), '4')
        bottom.set(qn('w:color'), '1e293b')
        pBdr.append(bottom)
        pPr.append(pBdr)

    def body(text, color=None, indent=False, bold=False, size=10):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(2)
        if indent:
            p.paragraph_format.left_indent = Inches(0.2)
        run = p.add_run(text)
        run.font.size = Pt(size)
        run.bold = bold
        run.font.color.rgb = color or mid
        return p

    def spacer():
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(6)

    # ── NAME ──
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run(data.full_name.upper())
    run.bold = True
    run.font.size = Pt(22)
    run.font.color.rgb = dark

    # ── HEADLINE ──
    p2 = doc.add_paragraph()
    p2.paragraph_format.space_after = Pt(4)
    run2 = p2.add_run(data.job_title or data.headline or "")
    run2.font.size = Pt(12)
    run2.font.color.rgb = mid

    # ── CONTACT ──
    contact_parts = [c for c in [data.email, data.phone, data.location, data.linkedin] if c]
    if contact_parts:
        body(" | ".join(contact_parts), color=mid, size=9)

    doc.add_paragraph()

    # ── PERSONAL INFO ──
    heading("Personal Information")
    personal_parts = [p for p in [
        f"Date of Birth: {data.date_of_birth}" if data.date_of_birth else "",
        f"Civil Status: {data.civil_status}" if data.civil_status else "",
        f"Nationality: {data.nationality}" if data.nationality else "",
        f"Availability: {data.availability}" if data.availability else "",
        f"Job Type: {data.job_type}" if data.job_type else "",
    ] if p]
    if personal_parts:
        body(" | ".join(personal_parts), size=9)
    else:
        spacer()

    # ── CAREER OBJECTIVE ──
    heading("Career Objective")
    # Will be filled by AI — leave space
    spacer()
    spacer()

    # ── WORK EXPERIENCE ──
    heading("Work Experience")
    if data.work_experience:
        for j in data.work_experience:
            if not (j.get("company") or j.get("role")): continue
            parts = []
            if j.get("role"): parts.append(j["role"])
            if j.get("company"): parts.append(j["company"])
            dur = f" | {j['duration']}" if j.get("duration") else ""
            body(f"{' — '.join(parts)}{dur}", color=dark, bold=True)
            for b in (j.get("responsibilities") or "").split("\n"):
                b = b.strip()
                if b:
                    body(f"• {b}", indent=True)
            spacer()
    else:
        spacer()
        spacer()

    # ── EDUCATION ──
    heading("Education")
    if data.education:
        for e in data.education:
            if not (e.get("school") or e.get("degree")): continue
            text = f"{e.get('degree', '')}"
            if e.get("school"): text += f" — {e['school']}"
            if e.get("year"): text += f" | {e['year']}"
            body(text, color=dark, bold=True)
    else:
        spacer()
        spacer()

    # ── KEY SKILLS ──
    heading("Key Skills")
    if data.skills.strip():
        skills = [s.strip() for s in data.skills.split(",") if s.strip()]
        body(", ".join(skills))
    else:
        spacer()
        spacer()

    # ── CERTIFICATIONS ──
    heading("Certifications & Licenses")
    if data.certifications:
        for c in data.certifications:
            if not c.get("name"): continue
            text = c["name"]
            if c.get("issuer"): text += f" — {c['issuer']}"
            if c.get("year"): text += f" ({c['year']})"
            body(f"• {text}", indent=True)
    else:
        spacer()
        spacer()

    # ── LANGUAGES ──
    heading("Languages")
    if data.languages:
        for l in data.languages:
            if not l.get("language"): continue
            text = l["language"]
            if l.get("proficiency"): text += f" ({l['proficiency']})"
            body(f"• {text}", indent=True)
    else:
        spacer()
        spacer()

    # ── SEMINARS ──
    heading("Seminars & Trainings")
    if data.seminars:
        for s in data.seminars:
            if not s.get("title"): continue
            text = s["title"]
            if s.get("organizer"): text += f" — {s['organizer']}"
            if s.get("year"): text += f" ({s['year']})"
            body(f"• {text}", indent=True)
    else:
        spacer()
        spacer()

    # ── REFERENCES ──
    heading("Character References")
    if data.references:
        for r in data.references:
            if not r.get("name"): continue
            text = " | ".join(filter(None, [r.get("name"), r.get("position"), r.get("company"), r.get("contact")]))
            body(f"• {text}", indent=True)
    else:
        body("Available upon request.", color=light)

    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)

    filename = (data.full_name or "resume").replace(" ", "_")
    return Response(
        content=buffer.getvalue(),
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": f"attachment; filename={filename}.docx"}
    )

@app.get("/health")
def health():
    return {"status": "ok"}